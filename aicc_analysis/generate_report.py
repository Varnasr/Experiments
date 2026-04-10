#!/usr/bin/env python3
"""
Generate DOCX report for AICC Observers DCC Readiness Grading.
"""

import json
import csv
import statistics
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ---- Helpers ----

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79", highlight_col=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            # Highlight grade column
            if highlight_col is not None and c_idx == highlight_col:
                grade = str(val).strip()
                colors_map = {"A": "C6EFCE", "B": "BDD7EE", "C": "FCE4D6", "D": "FFC7CE"}
                if grade in colors_map:
                    set_cell_shading(cell, colors_map[grade])

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def load_data():
    datasets = {}
    files = {
        "district_reports": "district_reports.json",
        "proposed_names": "proposed_names.json",
        "attachments": "attachments.json",
        "daily_reports": "daily_reports.json",
        "potential_leaders": "potential_leaders.json",
        "political_influencers": "political_influencers.json",
    }
    for key, fname in files.items():
        with open(fname) as f:
            datasets[key] = json.load(f)["Table"]
    return datasets


def load_grading_results():
    results = []
    with open("grading_results.csv") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row["FinalScore"] = float(row["FinalScore"])
            row["ProposedNames"] = int(row["ProposedNames"])
            row["DailyReports"] = int(row["DailyReports"])
            row["SupportingDocs"] = int(row["SupportingDocs"])
            row["PotentialLeaders"] = int(row["PotentialLeaders"])
            row["Influencers"] = int(row["Influencers"])
            row["DistrictReport"] = int(row["DistrictReport"])
            results.append(row)
    return results


def main():
    data = load_data()
    results = load_grading_results()
    results.sort(key=lambda x: x["FinalScore"], reverse=True)

    scores = [r["FinalScore"] for r in results]

    # Grade counts
    grade_counts = defaultdict(int)
    for r in results:
        grade_counts[r["Grade"]] += 1

    # State summaries
    state_grades = defaultdict(lambda: defaultdict(int))
    state_totals = defaultdict(int)
    state_scores = defaultdict(list)
    for r in results:
        state_grades[r["State"]][r["Grade"]] += 1
        state_totals[r["State"]] += 1
        state_scores[r["State"]].append(r["FinalScore"])

    # ============================================================
    # BUILD DOCUMENT
    # ============================================================
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # --- Title Page ---
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AICC OBSERVERS")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("DCC Readiness Grading Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph("")

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("A/B/C/D Grading System for Observer Performance\nin District Congress Committee Membership Preparation")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph("")
    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Total Districts Evaluated: {len(results)}\n")
    run.font.size = Pt(12)
    run = meta.add_run(f"States Covered: {len(state_totals)}\n")
    run.font.size = Pt(12)
    run = meta.add_run(f"Data Source: aiccobservers.in")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # --- TABLE OF CONTENTS (manual) ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Data Overview",
        "3. Grading Methodology",
        "    3.1 Parameters & Weights",
        "    3.2 Scoring Method",
        "    3.3 Grade Thresholds",
        "4. Results: Grade Distribution",
        "5. State-Wise Performance Summary",
        "6. Top 30 Performers (Grade A)",
        "7. Bottom 30 - Districts Needing Attention (Grade D)",
        "8. Detailed State-Wise Breakdown",
        "9. Recommendations",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # ============================================================
    # 1. EXECUTIVE SUMMARY
    # ============================================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        f"This report presents a comprehensive A/B/C/D grading assessment of AICC Observers "
        f"deployed across {len(state_totals)} states and {len(results)} districts to evaluate and prepare "
        f"candidates for District Congress Committee (DCC) and City Congress Committee (CCC) membership."
    )
    doc.add_paragraph(
        "The grading is based on six measurable parameters extracted from the aiccobservers.in platform, "
        "covering district profiling, candidate proposals, daily field activity, documentation quality, "
        "leadership identification, and stakeholder mapping."
    )

    doc.add_heading("Key Findings", level=2)
    bullets = [
        f"Grade A (Excellent): {grade_counts['A']} districts ({grade_counts['A']/len(results)*100:.1f}%) - Comprehensive work across all parameters.",
        f"Grade B (Good): {grade_counts['B']} districts ({grade_counts['B']/len(results)*100:.1f}%) - Substantial work with some gaps.",
        f"Grade C (Needs Improvement): {grade_counts['C']} districts ({grade_counts['C']/len(results)*100:.1f}%) - Significant gaps in coverage.",
        f"Grade D (Poor): {grade_counts['D']} districts ({grade_counts['D']/len(results)*100:.1f}%) - Minimal or no meaningful work done.",
        f"Average score across all districts: {statistics.mean(scores):.1f} / 100.",
        f"Best performing states: Odisha (69.6), Uttarakhand (67.9), Punjab (67.7), Himachal Pradesh (67.6).",
        f"States needing urgent attention: Maharashtra (16.5 avg, {state_grades['Maharashtra']['D']} Grade D), "
        f"Arunachal Pradesh (17.4 avg, {state_grades['Arunachal Pradesh']['D']} Grade D), Delhi (25.4 avg).",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()

    # ============================================================
    # 2. DATA OVERVIEW
    # ============================================================
    doc.add_heading("2. Data Overview", level=1)

    doc.add_paragraph(
        "Data was extracted from the AICC Observers platform (aiccobservers.in), which serves as the "
        "central reporting system for observers appointed to evaluate and prepare DCC/CCC membership across India. "
        "The platform collects the following categories of information from each observer per assigned district:"
    )

    data_sections = [
        ("General District Report", "District profile including number of assemblies, blocks, caste composition, "
         "local political issues, political factions, prominent Congress leaders, and current DCC President details."),
        ("Proposed Names for DCC/CCC", "Candidate profiles with detailed information: name, age, gender, caste, "
         "occupation, education, Congress party affiliation, positions held, election history, criminal record, "
         "strengths, weaknesses, and observer rating."),
        ("Daily Activity Reports", "Day-by-day reporting of field activities across 10 meeting types: "
         "Press Conferences, District/Assembly/Block/Booth-level meetings, Civil Society interactions, "
         "Group Representations, One-to-One interactions with aspirants and leaders."),
        ("Supporting Documents", "Uploaded evidence and documentation supporting the observer's assessments."),
        ("Potential Leaders Identified", "Leaders identified by the observer as having potential for party roles."),
        ("Non-Political Influencers", "Prominent non-political persons mapped for stakeholder engagement."),
    ]
    for title, desc in data_sections:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    # Data volume table
    doc.add_paragraph("")
    doc.add_heading("Data Volume Summary", level=2)
    add_styled_table(doc,
        ["Dataset", "Total Records", "Districts with Data", "Districts without Data"],
        [
            ["General District Reports", "595",
             str(sum(1 for d in data["district_reports"] if d.get("Profiles", 0) > 0)),
             str(sum(1 for d in data["district_reports"] if d.get("Profiles", 0) == 0))],
            ["Proposed Names", "595",
             str(sum(1 for d in data["proposed_names"] if d.get("Profiles", 0) > 0)),
             str(sum(1 for d in data["proposed_names"] if d.get("Profiles", 0) == 0))],
            ["Daily Reports", "532",
             str(sum(1 for d in data["daily_reports"] if (d.get("Total") or 0) > 0)),
             str(sum(1 for d in data["daily_reports"] if (d.get("Total") or 0) == 0))],
            ["Supporting Documents", "595",
             str(sum(1 for d in data["attachments"] if d.get("Profiles", 0) > 0)),
             str(sum(1 for d in data["attachments"] if d.get("Profiles", 0) == 0))],
            ["Potential Leaders", "595",
             str(sum(1 for d in data["potential_leaders"] if d.get("Profiles", 0) > 0)),
             str(sum(1 for d in data["potential_leaders"] if d.get("Profiles", 0) == 0))],
            ["Political Influencers", "595",
             str(sum(1 for d in data["political_influencers"] if d.get("Profiles", 0) > 0)),
             str(sum(1 for d in data["political_influencers"] if d.get("Profiles", 0) == 0))],
        ]
    )

    doc.add_page_break()

    # ============================================================
    # 3. GRADING METHODOLOGY
    # ============================================================
    doc.add_heading("3. Grading Methodology", level=1)

    doc.add_heading("3.1 Parameters & Weights", level=2)
    doc.add_paragraph(
        "Each district is scored on six parameters, weighted to reflect their relative importance "
        "in assessing DCC readiness:"
    )

    add_styled_table(doc,
        ["#", "Parameter", "Weight", "Rationale"],
        [
            ["1", "District Report Submitted", "15%", "Foundation: Did the observer complete the basic district profiling?"],
            ["2", "Proposed Names (DCC/CCC)", "25%", "Core Deliverable: Number of DCC/CCC candidates proposed with profiles."],
            ["3", "Daily Activity Reports", "25%", "Engagement: Frequency and consistency of field activities and meetings."],
            ["4", "Supporting Documents", "15%", "Evidence Quality: Documentation uploaded to support assessments."],
            ["5", "Potential Leaders Identified", "10%", "Depth: Going beyond mandatory work to identify future leaders."],
            ["6", "Non-Political Influencers", "10%", "Breadth: Mapping non-political stakeholders for outreach."],
        ],
        col_widths=[1, 5, 2, 9]
    )

    doc.add_heading("3.2 Scoring Method", level=2)
    doc.add_paragraph(
        "Each parameter is scored on a 0-100 scale using percentile ranking against all districts. "
        "This means a district's score on each parameter reflects how it compares to all other districts, "
        "not an absolute threshold. For example, if a district has more proposed names than 80% of all districts, "
        "it scores 80 on that parameter."
    )
    doc.add_paragraph(
        "The District Report parameter is binary (0 or 100) - either the report was submitted or it was not."
    )
    doc.add_paragraph(
        "The final score is the weighted sum of all six parameter scores."
    )

    doc.add_heading("3.3 Grade Thresholds", level=2)
    add_styled_table(doc,
        ["Grade", "Score Range", "Classification", "Description"],
        [
            ["A", ">= 70", "Excellent", "Comprehensive work done across all parameters. Observer has thoroughly covered district profiling, proposed adequate candidates, maintained active daily reporting, and uploaded supporting evidence."],
            ["B", "45 - 69", "Good", "Substantial work completed with some gaps. Most key deliverables met but room for improvement in one or more areas."],
            ["C", "20 - 44", "Needs Improvement", "Significant gaps in coverage. Core deliverables partially met. Observer needs to increase activity and documentation."],
            ["D", "< 20", "Poor", "Minimal or no meaningful work done. Critical deliverables missing. Immediate attention required."],
        ],
        col_widths=[1.5, 2.5, 3, 10],
        highlight_col=0
    )

    doc.add_page_break()

    # ============================================================
    # 4. RESULTS: GRADE DISTRIBUTION
    # ============================================================
    doc.add_heading("4. Results: Grade Distribution", level=1)

    add_styled_table(doc,
        ["Grade", "Districts", "Percentage", "Visual"],
        [
            ["A", str(grade_counts["A"]), f"{grade_counts['A']/len(results)*100:.1f}%",
             "|" + "#" * int(grade_counts["A"]/len(results)*50)],
            ["B", str(grade_counts["B"]), f"{grade_counts['B']/len(results)*100:.1f}%",
             "|" + "#" * int(grade_counts["B"]/len(results)*50)],
            ["C", str(grade_counts["C"]), f"{grade_counts['C']/len(results)*100:.1f}%",
             "|" + "#" * int(grade_counts["C"]/len(results)*50)],
            ["D", str(grade_counts["D"]), f"{grade_counts['D']/len(results)*100:.1f}%",
             "|" + "#" * int(grade_counts["D"]/len(results)*50)],
        ],
        highlight_col=0
    )

    doc.add_paragraph("")
    doc.add_heading("Score Statistics", level=2)
    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Total Districts Evaluated", str(len(results))],
            ["Mean Score", f"{statistics.mean(scores):.1f}"],
            ["Median Score", f"{statistics.median(scores):.1f}"],
            ["Standard Deviation", f"{statistics.stdev(scores):.1f}"],
            ["Minimum Score", f"{min(scores):.1f}"],
            ["Maximum Score", f"{max(scores):.1f}"],
        ]
    )

    doc.add_page_break()

    # ============================================================
    # 5. STATE-WISE PERFORMANCE
    # ============================================================
    doc.add_heading("5. State-Wise Performance Summary", level=1)

    state_rows = []
    for state in sorted(state_totals.keys()):
        sg = state_grades[state]
        avg = statistics.mean(state_scores[state])
        # Determine dominant grade
        dominant = max(["A", "B", "C", "D"], key=lambda g: sg[g])
        state_rows.append([
            state, str(state_totals[state]),
            str(sg["A"]), str(sg["B"]), str(sg["C"]), str(sg["D"]),
            f"{avg:.1f}", dominant
        ])

    # Sort by avg score descending
    state_rows.sort(key=lambda x: float(x[6]), reverse=True)

    add_styled_table(doc,
        ["State", "Total", "A", "B", "C", "D", "Avg Score", "Dominant"],
        state_rows,
        highlight_col=7
    )

    doc.add_page_break()

    # ============================================================
    # 6. TOP 30 PERFORMERS
    # ============================================================
    doc.add_heading("6. Top 30 Performers (Grade A)", level=1)

    top_rows = []
    for i, r in enumerate(results[:30], 1):
        top_rows.append([
            str(i), r["Observer"][:35], r["District"][:22], r["State"][:18],
            str(r["ProposedNames"]), str(r["DailyReports"]),
            str(r["SupportingDocs"]), f"{r['FinalScore']:.1f}", r["Grade"]
        ])

    add_styled_table(doc,
        ["#", "Observer", "District", "State", "Names", "Daily", "Docs", "Score", "Grade"],
        top_rows,
        highlight_col=8
    )

    doc.add_page_break()

    # ============================================================
    # 7. BOTTOM 30
    # ============================================================
    doc.add_heading("7. Bottom 30 - Districts Needing Attention", level=1)

    bottom_rows = []
    for i, r in enumerate(results[-30:], len(results) - 29):
        bottom_rows.append([
            str(i), r["Observer"][:35], r["District"][:22], r["State"][:18],
            str(r["ProposedNames"]), str(r["DailyReports"]),
            str(r["SupportingDocs"]), f"{r['FinalScore']:.1f}", r["Grade"]
        ])

    add_styled_table(doc,
        ["#", "Observer", "District", "State", "Names", "Daily", "Docs", "Score", "Grade"],
        bottom_rows,
        highlight_col=8
    )

    doc.add_page_break()

    # ============================================================
    # 8. STATE-WISE DETAILED BREAKDOWN
    # ============================================================
    doc.add_heading("8. Detailed State-Wise Breakdown", level=1)

    for state in sorted(state_totals.keys()):
        doc.add_heading(f"{state}", level=2)
        sg = state_grades[state]
        avg = statistics.mean(state_scores[state])

        p = doc.add_paragraph()
        run = p.add_run(f"Districts: {state_totals[state]}  |  ")
        run.font.size = Pt(10)
        run = p.add_run(f"Avg Score: {avg:.1f}  |  ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"A:{sg['A']}  B:{sg['B']}  C:{sg['C']}  D:{sg['D']}")
        run.font.size = Pt(10)

        state_results = [r for r in results if r["State"] == state]
        state_results.sort(key=lambda x: x["FinalScore"], reverse=True)

        s_rows = []
        for r in state_results:
            s_rows.append([
                r["District"][:25], r["Observer"][:30],
                str(r["DistrictReport"]), str(r["ProposedNames"]),
                str(r["DailyReports"]), str(r["SupportingDocs"]),
                str(r["PotentialLeaders"]),
                f"{r['FinalScore']:.1f}", r["Grade"]
            ])

        add_styled_table(doc,
            ["District", "Observer", "DR", "Names", "Daily", "Docs", "Leaders", "Score", "Grade"],
            s_rows,
            highlight_col=8
        )
        doc.add_paragraph("")

    doc.add_page_break()

    # ============================================================
    # 9. RECOMMENDATIONS
    # ============================================================
    doc.add_heading("9. Recommendations", level=1)

    doc.add_heading("Immediate Actions (Grade D Districts)", level=2)
    recs_d = [
        f"92 districts (15.5%) scored Grade D with minimal or no work. Maharashtra alone accounts for 60 of these.",
        "Observers in Maharashtra, Arunachal Pradesh, and Delhi need immediate engagement and re-briefing on deliverables.",
        "Consider reassigning observers who have shown zero activity across all parameters.",
        "Set a 2-week deadline for Grade D observers to submit at minimum: District Report + 3 Proposed Names.",
    ]
    for r in recs_d:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("Improvement Areas (Grade C Districts)", level=2)
    recs_c = [
        "86 districts (14.5%) scored Grade C - partial work done but significant gaps remain.",
        "Focus on increasing Daily Reporting frequency - many Grade C districts have low activity counts.",
        "Encourage upload of supporting documents - only 300 of 595 districts have any documentation.",
        "Political Influencer mapping is the weakest parameter overall - only 50 districts have any data.",
    ]
    for r in recs_c:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("Best Practices (from Grade A Districts)", level=2)
    recs_a = [
        "Top performers combine high daily reporting frequency with thorough proposed name profiles.",
        "States like Odisha, Jharkhand, and Punjab show consistently high performance across most districts.",
        "Observers handling multiple districts (e.g., SHRI BISHWARANJAN MOHANTHY across 3 states) can maintain Grade A - workload is not necessarily a barrier.",
        "Consider recognizing and rewarding top-performing observers to incentivize continued excellence.",
    ]
    for r in recs_a:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("System Improvements", level=2)
    recs_sys = [
        "Make the grading dashboard visible to observers in real-time to drive self-improvement.",
        "Add mandatory minimum thresholds: no submission should be accepted without at least a District Report.",
        "Integrate weekly automated grading emails to state-level supervisors.",
        "Consider adding a 'Quality Score' for Proposed Names based on completeness of candidate profiles.",
    ]
    for r in recs_sys:
        doc.add_paragraph(r, style="List Bullet")

    # ============================================================
    # SAVE
    # ============================================================
    output_file = "AICC_Observers_DCC_Grading_Report.docx"
    doc.save(output_file)
    print(f"Report saved: {output_file}")


if __name__ == "__main__":
    main()
